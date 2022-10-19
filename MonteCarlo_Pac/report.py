from docx import Document
from MonteCarlo_Pac import cashflows
import numpy as np
import datetime

def report(distribution_chosen,mc_simulation,csh_flw,ticker,total_dist_no_hedging,total_dist_with_hedging):
	document = Document()

	document.add_heading('The distribution of CashFlow values during the year with and without hedging',level = 0)

	document.add_paragraph('{} distribution was selected for Monte Carlo simulation of asset price motions during the year. '
		.format({1: 'LogNormal',2: 'LogNormal with Jump',3:'LogLaplace',4: 'LogLaplace with Jump',5: 'Hypersecant',6: 'Cauchy'}[distribution_chosen]))
	document.add_paragraph('{} motions were simulated in total and {} CashFlows, that are subject to risk of adverse {} movements, were added:'
		.format(len(mc_simulation[0]),len(csh_flw),ticker))

	#Table 1(Cashflow list)
	document.add_heading('All Cashflows',level = 1)
	table1 = document.add_table(rows=1+len(csh_flw), cols=4)
	table1.cell(0,0).text = 'Date'
	table1.cell(0,1).text = 'Ammount'
	table1.cell(0,2).text = 'Buy'
	table1.cell(0,3).text = 'Hedging strategy'
	for i in range(1,1+len(csh_flw)):
		table1.cell(i,0).text = str(csh_flw[i-1].date)
		table1.cell(i,1).text = str(csh_flw[i-1].ammount)
		table1.cell(i,2).text = str(csh_flw[i-1].buy)
		table1.cell(i,3).text = str(csh_flw[i-1].hed_strat)

	document.add_page_break()

	document.add_heading('No hedging',level = 1)
	document.add_paragraph('Given the distribution of {} and the future cashflows, risks for the enterprise are the following:'
		.format(ticker))


	#Table 2(No hedging)
	#Values at risk
	total_dist_no_hedging.sort()
	var5 = total_dist_no_hedging[int(len(total_dist_no_hedging)*0.05)]
	var5 = format(var5,'.4f')
	var1 = total_dist_no_hedging[int(len(total_dist_no_hedging)*0.01)]
	var1 = format(var1,'.4f')

	table2 = document.add_table(rows=4, cols=2)
	table2.cell(0,0).text = 'The worst total value of CFs '
	table2.cell(0,1).text = str(format(min(total_dist_no_hedging),'.4f'))

	table2.cell(1,0).text = 'The best total value of CFs '
	table2.cell(1,1).text = str(format(max(total_dist_no_hedging),'.4f'))

	table2.cell(2,0).text = 'VaR 1%'
	table2.cell(2,1).text = str(var1)

	table2.cell(3,0).text = 'VaR 5%'
	table2.cell(3,1).text = str(var5)

	#Plot distribution for better wisualization
	document.add_heading('Plot of total CashFlow value distribution WITHOUT HEDGING',level = 2)
	document.add_picture('hedging_absence_distribution.png')

	document.add_page_break()


	document.add_heading('Hedging implemented',level = 1)

	#Table 3 with hedging
	total_dist_with_hedging.sort()
	var5 = total_dist_with_hedging[int(len(total_dist_with_hedging)*0.05)]
	var5 = format(var5,'.4f')
	var1 = total_dist_with_hedging[int(len(total_dist_with_hedging)*0.01)]
	var1 = format(var1,'.4f')

	table3 = document.add_table(rows=4, cols=2)
	table3.cell(0,0).text = 'The worst total value of CFs'
	table3.cell(0,1).text = str(format(min(total_dist_with_hedging),'.4f'))

	table3.cell(1,0).text = 'The best total value of CFs'
	table3.cell(1,1).text = str(format(max(total_dist_with_hedging),'.4f'))

	table3.cell(2,0).text = 'VaR 1%'
	table3.cell(2,1).text = str(var1)

	table3.cell(3,0).text = 'VaR 5%'
	table3.cell(3,1).text = str(var5)

	#Plot distribution for better wisualization
	document.add_heading('Plot of total CashFlow value distribution WITH HEDGING',level = 2)
	document.add_picture('hedging_presence_distribution.png')

	#Table 4: effect of the hedging strategy
	table4 = document.add_table(rows=4, cols=2)

	min_value = min(total_dist_with_hedging)
	bad_cases_avoided = 0

	for i in total_dist_no_hedging:
		if i < min_value:
			bad_cases_avoided += 1
		else:
			pass

	max_value = max(total_dist_with_hedging)
	good_cases_avoided = 0

	for i in total_dist_no_hedging:
		if i > max_value:
			good_cases_avoided += 1
		else:
			pass

	table4.cell(0,0).text = 'Proportion of bad cases avoided'
	table4.cell(0,1).text = str( 100 * bad_cases_avoided / len(total_dist_with_hedging) ) + '%'

	table4.cell(1,0).text = 'Proportion of good cases avoided'
	table4.cell(1,1).text = str( 100 * good_cases_avoided / len(total_dist_with_hedging) ) + '%'

	table4.cell(2,0).text = 'Worst case difference'
	table4.cell(2,1).text = str( abs(min_value - min(total_dist_no_hedging)) )

	table4.cell(3,0).text = 'Best case difference'
	table4.cell(3,1).text = str( abs(max_value - min(total_dist_no_hedging)) )


	document.save('report.docx')